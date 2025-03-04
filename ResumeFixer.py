import tkinter as tk
from tkinter import filedialog, messagebox
import requests
from bs4 import BeautifulSoup
import openai
from docx import Document
from pswd import chatgptkey

# OpenAI API Key
API_KEY = chatgptkey
openai.api_key = API_KEY

class JobApplicationGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Job Application Assistant")

        # URL Input
        self.url_label = tk.Label(root, text="Enter the job listing URL:")
        self.url_label.pack(pady=5)

        self.url_entry = tk.Entry(root, width=50)
        self.url_entry.pack(pady=5)

        self.fetch_button = tk.Button(root, text="Fetch Job Info", command=self.fetch_job_info)
        self.fetch_button.pack(pady=5)

        # Job Summary Display
        self.summary_text = tk.Text(root, height=10, width=60)
        self.summary_text.pack(pady=5)
        self.summary_text.config(state=tk.DISABLED)

        # Resume Upload Section
        self.upload_button = tk.Button(root, text="Upload Resume", command=self.upload_resume)
        self.upload_button.pack(pady=5)

        # Resume Display Section
        self.resume_text = tk.Text(root, height=15, width=60)
        self.resume_text.pack(pady=5)
        self.resume_text.config(state=tk.DISABLED)

        self.resume_path = None
        self.resume_doc = None  # Store Document object

        # Edit and Save Buttons
        self.edit_button = tk.Button(root, text="Edit Resume", command=self.edit_resume, state=tk.DISABLED)
        self.edit_button.pack(pady=5)

        self.save_button = tk.Button(root, text="Download Edited Resume", command=self.save_resume, state=tk.DISABLED)
        self.save_button.pack(pady=5)

    def fetch_job_info(self):
        url = self.url_entry.get().strip()
        if not url:
            messagebox.showwarning("Input Error", "Please provide a URL.")
            return

        try:
            job_content = self.get_webpage_content(url)
            job_summary = self.get_job_summary(job_content)
            self.display_job_summary(job_summary)
        except Exception as e:
            messagebox.showerror("Error", f"Error fetching or processing the job listing: {e}")

    def get_webpage_content(self, url):
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        return soup.get_text(separator=" ", strip=True)

    def get_job_summary(self, content):
        try:
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": f"Summarize the following job listing:\n\n{content}"}
                ]
            )
            return response['choices'][0]['message']['content'].strip()
        except Exception as e:
            raise Exception(f"Error with GPT API: {e}")

    def display_job_summary(self, summary):
        self.summary_text.config(state=tk.NORMAL)
        self.summary_text.delete(1.0, tk.END)
        self.summary_text.insert(tk.END, summary)
        self.summary_text.config(state=tk.DISABLED)
        self.edit_button.config(state=tk.NORMAL)

    def upload_resume(self):
        file_path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if file_path:
            self.resume_path = file_path
            self.resume_doc = Document(file_path)  # Store Document object
            self.display_resume()
            self.save_button.config(state=tk.NORMAL)

    def display_resume(self):
        self.resume_text.config(state=tk.NORMAL)
        self.resume_text.delete(1.0, tk.END)
        for para in self.resume_doc.paragraphs:
            self.resume_text.insert(tk.END, para.text + '\n')
        self.resume_text.config(state=tk.DISABLED)

    def edit_resume(self):
        if not self.resume_doc or not self.resume_path:
            messagebox.showwarning("No Resume", "Please upload a resume first.")
            return

        job_summary = self.summary_text.get(1.0, tk.END).strip()
        if not job_summary:
            messagebox.showwarning("No Job Summary", "Please fetch the job summary first.")
            return

        edited_paragraphs = self.get_edited_resume([para.text for para in self.resume_doc.paragraphs], job_summary)
        for para, edited_text in zip(self.resume_doc.paragraphs, edited_paragraphs):
            para.text = edited_text

        self.display_resume()

    def get_edited_resume(self, resume_content, job_summary):
        try:
            resume_text = "\n".join(resume_content)
            prompt = f"Edit the following resume to match the job requirements:\nJob Summary: {job_summary}\n\nResume:\n{resume_text}"
            response = openai.ChatCompletion.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt}
                ]
            )
            return response['choices'][0]['message']['content'].strip().split("\n")
        except Exception as e:
            messagebox.showerror("Error", f"Error editing resume: {e}")
            return resume_content

    def save_resume(self):
        if not self.resume_doc:
            messagebox.showwarning("No Content", "There is no resume content to save.")
            return

        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Documents", "*.docx")])
        if save_path:
            self.resume_doc.save(save_path)
            messagebox.showinfo("Success", f"Edited resume saved to {save_path}")

# Run GUI
root = tk.Tk()
app = JobApplicationGUI(root)
root.mainloop()
